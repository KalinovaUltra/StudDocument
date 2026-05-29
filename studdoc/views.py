from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse
from django.views.decorators.http import require_POST
import json
from docx import Document
import re
from django.contrib.auth.models import User
from .forms import RequestForm
from .models import DocumentTemplate, RequestFile, UserProfile, Request
from docxtpl import DocxTemplate
from django.http import HttpResponse
from io import BytesIO
import zipfile

# ===== ГЛАВНАЯ СТРАНИЦА =====
def index(request):
    if request.method == 'POST' and request.user.is_authenticated:
        form = RequestForm(request.POST, request.FILES)
        if form.is_valid():
            request_obj = form.save(commit=False)
            request_obj.student = request.user
            template = DocumentTemplate.objects.filter(category=form.cleaned_data['category']).first()
            if template:
                request_obj.template = template
            request_obj.save()
            for f in request.FILES.getlist('attachments'):
                RequestFile.objects.create(request=request_obj, file=f)
            messages.success(request, '✅ Заявка успешно отправлена!')
            return redirect('index')
    else:
        form = RequestForm()
    return render(request, 'studdoc/index.html', {'form': form})

# ===== ЛИЧНЫЙ КАБИНЕТ СТУДЕНТА =====
@login_required
def profile(request):
    user_profile, _ = UserProfile.objects.get_or_create(user=request.user)
    user_requests = Request.objects.filter(student=request.user).order_by('-created_at')

    requests_data = []
    for req in user_requests:
        requests_data.append({
            'id': req.id,
            'title': req.get_category_display(),
            'category': req.get_category_display(),
            'comment': req.comment or '',
            'date': req.created_at.strftime('%d.%m.%Y %H:%M'),
            'status': req.get_status_display(),
            'statusClass': {
                'pending': 'status-pending',
                'in_profcom': 'status-in-profcom',
                'profcom_approved': 'status-profcom-approved',
                'profcom_rejected': 'status-profcom-rejected',
                'approved': 'status-approved',
                'rejected': 'status-rejected',
            }.get(req.status, 'status-pending'),
            'has_files': req.attached_files.exists()
        })

    context = {
        'user_profile': user_profile,
        'requests_json': json.dumps(requests_data, ensure_ascii=False),
        'page_title': 'Личный кабинет'
    }
    return render(request, 'studdoc/profile.html', context)

# ===== API: Обновление профиля =====
@require_POST
@login_required
def update_profile(request):
    try:
        data = json.loads(request.body)
        profile = request.user.profile
        user = request.user

        profile.last_name = data.get('last_name', '')
        profile.first_name = data.get('first_name', '')
        profile.patronymic = data.get('patronymic', '')
        profile.group = data.get('group', '')
        profile.faculty = data.get('faculty', '')
        profile.save()

        new_username = data.get('username', '')
        if new_username and new_username != user.username:
            if User.objects.exclude(id=user.id).filter(username=new_username).exists():
                return JsonResponse({'success': False, 'message': 'Логин уже занят'}, status=400)
            user.username = new_username

        new_email = data.get('email', '')
        if new_email and new_email != user.email:
            if User.objects.exclude(id=user.id).filter(email=new_email).exists():
                return JsonResponse({'success': False, 'message': 'Email уже используется'}, status=400)
            user.email = new_email

        user.save()

        return JsonResponse({
            'success': True,
            'full_name': profile.get_full_name(),
            'username': user.username,
            'email': user.email,
            'group': profile.group,
            'faculty': profile.faculty
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=400)

# ===== API: Создание заявки (единственная версия) =====
@require_POST
@login_required
def create_request_api(request):
    try:
        category = request.POST.get('category')
        comment = request.POST.get('comment', '')
        if not category:
            return JsonResponse({'error': 'Не выбрана категория'}, status=400)

        template_obj = DocumentTemplate.objects.filter(category=category).first()
        request_obj = Request.objects.create(
            student=request.user,
            category=category,
            comment=comment,
            status='pending',
            template=template_obj
        )

        template_data = {}
        for key, value in request.POST.items():
            if key not in ['csrfmiddlewaretoken', 'category', 'comment']:
                template_data[key] = value
        request_obj.form_data = json.dumps(template_data, ensure_ascii=False)
        request_obj.save()

        for f in request.FILES.getlist('attachments'):
            RequestFile.objects.create(request=request_obj, file=f)

        return JsonResponse({'success': True, 'message': 'Заявка создана', 'request_id': request_obj.id})
    except Exception as e:
        print("ОШИБКА:", e)
        return JsonResponse({'error': str(e)}, status=500)

# ===== API: Детали заявки (для студента) =====
@login_required
def request_detail(request, request_id):
    req = get_object_or_404(Request, id=request_id, student=request.user)

    files = [{'name': f.file.name.split('/')[-1], 'url': f.file.url} for f in req.attached_files.all()]

    data = {
        'id': req.id,
        'title': req.get_category_display(),
        'category': req.category,
        'comment': req.comment or 'Нет комментария',
        'date': req.created_at.strftime('%d.%m.%Y %H:%M'),
        'status': req.get_status_display(),
        'statusClass': {
            'pending': 'status-pending',
            'approved': 'status-approved',
            'rejected': 'status-rejected'
        }.get(req.status, 'status-pending'),
        'files': files
    }
    return JsonResponse(data)

# ===== API: Скачивание WORD-документа =====
@login_required
def download_request_docx(request, request_id):
    req = get_object_or_404(Request, id=request_id, student=request.user)
    template_obj = req.template
    if not template_obj or not template_obj.file:
        return HttpResponse("Шаблон не найден", status=404)

    doc = DocxTemplate(template_obj.file.path)
    form_data = json.loads(req.form_data) if req.form_data else {}

    context = {}
    if 'name' in form_data:
        context['full_name'] = form_data['name']
    if 'group' in form_data:
        context['group'] = form_data['group']
    if 'faculty' in form_data:
        context['faculty'] = form_data['faculty']
    if 'date' in form_data:
        context['date'] = form_data['date']
    if 'reason' in form_data:
        context['reason'] = form_data['reason']
    if 'comment' in form_data:
        context['comment'] = form_data['comment']
    for key, value in form_data.items():
        if key not in ['name', 'group', 'faculty', 'date', 'reason', 'comment']:
            context[key] = value

    doc.render(context)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    response = HttpResponse(
        file_stream,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="request_{req.id}.docx"'
    return response

# ===== API: Предпросмотр шаблона (при заполнении формы) =====
@login_required
def preview_template_api(request, category):
    if request.method != 'POST':
        return JsonResponse({'error': 'Метод не разрешён'}, status=405)

    try:
        data = json.loads(request.body)
        values = data.get('values', {})
        context = {}
        if 'name' in values:
            context['full_name'] = values['name']
        if 'group' in values:
            context['group'] = values['group']
        if 'faculty' in values:
            context['faculty'] = values['faculty']
        if 'reason' in values:
            context['reason'] = values['reason']
        if 'comment' in values:
            context['comment'] = values['comment']
        for key, value in values.items():
            if key not in ['name', 'group', 'faculty', 'reason', 'comment']:
                context[key] = value

        template_obj = DocumentTemplate.objects.filter(category=category).first()
        if not template_obj or not template_obj.file:
            return JsonResponse({'error': 'Шаблон не найден'}, status=404)

        doc = Document(template_obj.file.path)
        full_text = [para.text for para in doc.paragraphs]
        text = '\n'.join(full_text)

        def replace_placeholder(match):
            key = match.group(1).strip()
            return context.get(key, '')

        html_text = re.sub(r'{{(.*?)}}', replace_placeholder, text)
        html_text = html_text.replace('\n', '<br>')

        return JsonResponse({'success': True, 'html': f'<div class="preview-document-text">{html_text}</div>'})
    except Exception as e:
        print("Ошибка в preview_template_api:", str(e))
        return JsonResponse({'error': str(e)}, status=500)

# ===== API: Превью для модального окна =====
@login_required
def request_preview_html(request, request_id):
    req = get_object_or_404(Request, id=request_id)

    allowed_roles = ['staff', 'admin', 'profcom']
    is_allowed_role = request.user.profile.role in allowed_roles
    is_owner = req.student == request.user
    if not (is_owner or is_allowed_role):
        return JsonResponse({'error': 'Нет прав'}, status=403)

    template_obj = req.template
    if not template_obj or not template_obj.file:
        return JsonResponse({'error': 'Шаблон не найден'}, status=404)

    form_data = json.loads(req.form_data) if req.form_data else {}
    context = {}
    if 'name' in form_data:
        context['full_name'] = form_data['name']
    if 'group' in form_data:
        context['group'] = form_data['group']
    if 'faculty' in form_data:
        context['faculty'] = form_data['faculty']
    if 'reason' in form_data:
        context['reason'] = form_data['reason']
    if 'comment' in form_data:
        context['comment'] = form_data['comment']
    if 'date' in form_data:
        context['date'] = form_data['date']
    for key, value in form_data.items():
        if key not in context:
            context[key] = value

    from docxtpl import DocxTemplate
    from django.utils.html import escape

    try:
        doc = DocxTemplate(template_obj.file.path)
        doc.render(context)

        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                escaped = escape(text)
                full_text.append(escaped.replace('\n', '<br>'))
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_text.append(escape(cell_text))
                if row_text:
                    full_text.append('<br>'.join(row_text))
        rendered_text = '<br><br>'.join(full_text) if full_text else '<em>Документ пуст</em>'
    except Exception as e:
        print(f"❌ Ошибка рендеринга превью: {e}")
        rendered_text = f'<p class="error">Ошибка загрузки: {escape(str(e))}</p>'

    status_display = req.get_status_display()
    date_str = req.created_at.strftime('%d.%m.%Y %H:%M')
    category_display = req.get_category_display()
    comment_text = escape(req.comment) if req.comment else '<em>нет</em>'

    full_html = f"""
    <div class="modal-document-header">
        <h3>{escape(category_display)}</h3>
        <p><strong>Категория:</strong> {escape(category_display)}</p>
        <p><strong>Статус:</strong> {escape(status_display)}</p>
        <p><strong>Дата подачи:</strong> {date_str}</p>
        <p><strong>Комментарий:</strong> {comment_text}</p>
    </div>
    <div class="modal-document-body" style="white-space: pre-wrap; line-height: 1.6;">
        {rendered_text}
    </div>
    <div class="modal-document-footer">
        <a href="/download-request/{req.id}/" class="btn-download" target="_blank">📄 Скачать doc</a>
    </div>
    """
    return JsonResponse({'success': True, 'html': full_html})

# ===== ПАНЕЛЬ ДЕКАНАТА =====
@login_required
def deanery_panel(request):
    if request.user.profile.role not in ['staff', 'admin']:
        messages.error(request, 'У вас нет доступа к этой странице.')
        return redirect('profile')

    all_requests = Request.objects.all().select_related('student', 'student__profile').order_by('-created_at')
    requests_data = []
    for req in all_requests:
        requests_data.append({
            'id': req.id,
            'student_name': req.student.profile.get_full_name() or req.student.username,
            'student_username': req.student.username,
            'title': req.get_category_display(),
            'comment': req.comment or '',
            'date': req.created_at.strftime('%d.%m.%Y %H:%M'),
            'status': req.status,
            'status_display': req.get_status_display(),
            'statusClass': {
                'pending': 'status-pending',
                'in_profcom': 'status-in-profcom',
                'profcom_approved': 'status-profcom-approved',
                'profcom_rejected': 'status-profcom-rejected',
                'approved': 'status-approved',
                'rejected': 'status-rejected',
            }.get(req.status, 'status-pending'),
            'has_files': req.attached_files.exists()
        })
    return render(request, 'studdoc/deanery_panel.html', {
        'requests_json': json.dumps(requests_data, ensure_ascii=False)
    })

# ===== API: Обновление статуса =====
@require_POST
@login_required
def update_request_status(request, request_id):
    role = request.user.profile.role
    if role not in ['staff', 'admin', 'profcom']:
        return JsonResponse({'error': 'Нет прав'}, status=403)

    try:
        data = json.loads(request.body)
        new_status = data.get('status')
        allowed_for_staff = ['pending', 'in_profcom', 'approved', 'rejected']
        allowed_for_profcom = ['in_profcom', 'profcom_approved', 'profcom_rejected']
        if role == 'staff' and new_status not in allowed_for_staff:
            return JsonResponse({'error': 'Некорректный статус для деканата'}, status=400)
        if role == 'profcom' and new_status not in allowed_for_profcom:
            return JsonResponse({'error': 'Некорректный статус для профкома'}, status=400)

        req = get_object_or_404(Request, id=request_id)
        req.status = new_status
        req.save()

        return JsonResponse({
            'success': True,
            'new_status': req.get_status_display(),
            'statusClass': {
                'pending': 'status-pending',
                'in_profcom': 'status-in-profcom',
                'profcom_approved': 'status-profcom-approved',
                'profcom_rejected': 'status-profcom-rejected',
                'approved': 'status-approved',
                'rejected': 'status-rejected',
            }.get(req.status, 'status-pending')
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# ===== API: Детали заявки для деканата =====
@login_required
def deanery_request_detail(request, request_id):
    if request.user.profile.role not in ['staff', 'admin']:
        return JsonResponse({'error': 'Нет прав'}, status=403)

    req = get_object_or_404(Request, id=request_id)
    files = [{'name': f.file.name.split('/')[-1], 'url': f.file.url} for f in req.attached_files.all()]

    data = {
        'id': req.id,
        'title': req.get_category_display(),
        'category': req.category,
        'comment': req.comment or 'Нет комментария',
        'date': req.created_at.strftime('%d.%m.%Y %H:%M'),
        'status': req.get_status_display(),
        'statusClass': {
            'pending': 'status-pending',
            'approved': 'status-approved',
            'rejected': 'status-rejected'
        }.get(req.status, 'status-pending'),
        'files': files
    }
    return JsonResponse(data)

# ===== ПАНЕЛЬ ПРОФКОМА =====
@login_required
def profcom_panel(request):
    if request.user.profile.role not in ['profcom', 'admin']:
        messages.error(request, 'У вас нет доступа к этой странице.')
        return redirect('profile')

    profcom_requests = Request.objects.filter(
        status__in=['in_profcom', 'profcom_approved', 'profcom_rejected']
    ).select_related('student', 'student__profile').order_by('-created_at')

    requests_data = []
    for req in profcom_requests:
        requests_data.append({
            'id': req.id,
            'student_name': req.student.profile.get_full_name() or req.student.username,
            'student_username': req.student.username,
            'title': req.get_category_display(),
            'comment': req.comment or '',
            'date': req.created_at.strftime('%d.%m.%Y %H:%M'),
            'status': req.status,
            'status_display': req.get_status_display(),
            'statusClass': {
                'pending': 'status-pending',
                'in_profcom': 'status-in-profcom',
                'profcom_approved': 'status-profcom-approved',
                'profcom_rejected': 'status-profcom-rejected',
                'approved': 'status-approved',
                'rejected': 'status-rejected',
            }.get(req.status, 'status-pending'),
            'has_files': req.attached_files.exists()
        })
    return render(request, 'studdoc/profcom_panel.html', {
        'requests_json': json.dumps(requests_data, ensure_ascii=False)
    })

# ===== API: Скачать прикреплённые файлы (ZIP) =====
@login_required
def download_attachments(request, request_id):
    req = get_object_or_404(Request, id=request_id)

    is_owner = req.student == request.user
    role = request.user.profile.role
    if not (is_owner or role in ['staff', 'admin', 'profcom']):
        return JsonResponse({'error': 'Нет прав'}, status=403)

    attached_files = req.attached_files.all()
    if not attached_files:
        return JsonResponse({'error': 'Нет прикреплённых файлов'}, status=404)

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for af in attached_files:
            file_path = af.file.path
            arcname = af.file.name.split('/')[-1]
            zip_file.write(file_path, arcname)

    zip_buffer.seek(0)
    response = HttpResponse(zip_buffer, content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="attachments_{request_id}.zip"'
    return response