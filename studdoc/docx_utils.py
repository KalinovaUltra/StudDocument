# studdoc/docx_utils.py

import re
from io import BytesIO
from docxtpl import DocxTemplate
from django.core.files.base import ContentFile

def extract_placeholders(docx_file):
    """
    Извлекает все имена полей из DOCX-шаблона.
    """
    doc = DocxTemplate(docx_file)
    from docx import Document
    document = Document(docx_file)
    placeholders = set()
    pattern = r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}'
    for paragraph in document.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        placeholders.update(matches)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                placeholders.update(matches)
    return list(placeholders)

def generate_docx_from_template(template_file, context):
    """
    Принимает файл-шаблон.
    Возвращает BytesIO с готовым DOCX.
    """
    doc = DocxTemplate(template_file)
    doc.render(context)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer